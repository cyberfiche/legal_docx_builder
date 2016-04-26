from docx import Document
import glob

# use glob to create an iterator that will cycle through the docx files in a directory
docpaths = (name for name in glob.glob('/path/*.docx'))
masterDoc = Document()

#iterate through the file names in docpaths
for docpath in docpaths:
    #create a python-docx object from the current file name
    doc = Document(docpath)
    #iterate through the paragraphs in doc and copy the text and paragraph styles to masterDoc
    for para in doc.paragraphs:
        masterDoc.add_paragraph(text=para.text, style=para.style)

#save masterDoc to disk
masterDoc.save('/path/masterDoc.docx')
