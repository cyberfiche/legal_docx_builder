from docx import Document

masterDoc = Document()
p = 0
t = 0

for elem in elems:
    if elem.tag.endswith('p'):
        masterDoc.add_paragraph(doc.paragraphs[p].text, doc.paragraphs[p].style)
        p += 1
    if elem.tag.endswith('tbl'):
        tbl = masterDoc.add_table(len(doc.tables[t].rows),len(doc.tables[t].columns), doc.tables[t].style)
        #for i in range(0, len(tbl.rows)):
        tbl = doc.tables[t] ##This does not correctly copy the contents of the table
